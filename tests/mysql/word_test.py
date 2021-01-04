# -*- coding: utf8 -*-
import unittest
import json

from docx import Document

from django.test import TestCase
from www.models.main import Tenants
from www.models.main import TenantRegionInfo
from www.models.main import TenantRegionResource
from www.models.main import ServiceInfo
from www.models.main import TenantServiceInfo
from www.models.main import TenantServiceInfoDelete
from www.models.main import TenantServiceLog
from www.models.main import TenantServiceRelation
from www.models.main import TenantServiceEnv
from www.models.main import TenantServiceAuth
from www.models.main import TenantServiceExtendMethod
from www.models.main import ServiceDomain
from www.models.main import ServiceDomainCertificate
from www.models.main import PermRelService
from www.models.main import PermRelTenant
from www.models.main import PhoneCode
from www.models.main import TenantServiceL7Info
from www.models.main import TenantServiceEnvVar
from www.models.main import TenantServicesPort
from www.models.main import TenantServiceMountRelation
from www.models.main import TenantServiceVolume
from www.models.main import TenantServiceConfigurationFile
from www.models.main import ServiceGroup
from www.models.main import ServiceGroupRelation
from www.models.main import ImageServiceRelation
from www.models.main import ComposeServiceRelation
from www.models.main import ServiceRule
from www.models.main import ServiceRuleHistory
from www.models.main import ServiceCreateStep
from www.models.main import ServiceProbe
from www.models.main import ConsoleConfig
from www.models.main import TenantEnterprise
from www.models.main import TenantEnterpriseToken
from www.models.main import TenantServiceGroup
from www.models.main import ServiceTcpDomain
from www.models.main import ThirdPartyServiceEndpoints
from www.models.main import ServiceWebhooks
from www.models.main import GatewayCustomConfiguration

from console.models.main import ConsoleSysConfig
from console.models.main import RainbondCenterApp
from console.models.main import RainbondCenterAppInherit
from console.models.main import RainbondCenterPlugin
from console.models.main import ServiceShareRecord
from console.models.main import EnterpriseUserPerm
from console.models.main import TenantUserRole
from console.models.main import TenantUserPermission
from console.models.main import TenantUserRolePermission
from console.models.main import PermGroup
from console.models.main import ServiceRelPerms
from console.models.main import AppExportRecord
from console.models.main import UserMessage
from console.models.main import AppImportRecord
from console.models.main import GroupAppBackupRecord
from console.models.main import GroupAppMigrateRecord
from console.models.main import GroupAppBackupImportRecord
from console.models.main import Applicants
from console.models.main import DeployRelation
from console.models.main import ServiceBuildSource
from console.models.main import TenantServiceBackup
from console.models.main import AppUpgradeRecord
from console.models.main import ServiceUpgradeRecord
from console.models.main import RegionConfig
from console.models.main import CloundBangImages
from console.models.main import Announcement


class Json2WordCase(TestCase):
    def init(self):
        self.sources = [
            Tenants(),
            TenantRegionInfo(),
            TenantRegionResource(),
            ServiceInfo(),
            TenantServiceInfo(),
            TenantServiceInfoDelete(),
            TenantServiceLog(),
            TenantServiceRelation(),
            TenantServiceEnv(),
            TenantServiceAuth(),
            TenantServiceExtendMethod(),
            ServiceDomain(),
            ServiceDomainCertificate(),
            PermRelService(),
            PermRelTenant(),
            PhoneCode(),
            TenantServiceL7Info(),
            TenantServiceEnvVar(),
            TenantServicesPort(),
            TenantServiceMountRelation(),
            TenantServiceVolume(),
            TenantServiceConfigurationFile(),
            ServiceGroup(),
            ServiceGroupRelation(),
            ImageServiceRelation(),
            ComposeServiceRelation(),
            ServiceRule(),
            ServiceRuleHistory(),
            ServiceCreateStep(),
            ServiceProbe(),
            ConsoleConfig(),
            TenantEnterprise(),
            TenantEnterpriseToken(),
            TenantServiceGroup(),
            ServiceTcpDomain(),
            ThirdPartyServiceEndpoints(),
            ServiceWebhooks(),
            GatewayCustomConfiguration(),
            ConsoleSysConfig(),
            RainbondCenterApp(),
            RainbondCenterAppInherit(),
            RainbondCenterPlugin(),
            ServiceShareRecord(),
            EnterpriseUserPerm(),
            TenantUserRole(),
            TenantUserPermission(),
            TenantUserRolePermission(),
            PermGroup(),
            ServiceRelPerms(),
            AppExportRecord(),
            UserMessage(),
            AppImportRecord(),
            GroupAppBackupRecord(),
            GroupAppMigrateRecord(),
            GroupAppBackupImportRecord(),
            Applicants(),
            DeployRelation(),
            ServiceBuildSource(),
            TenantServiceBackup(),
            AppUpgradeRecord(),
            ServiceUpgradeRecord(),
            RegionConfig(),
            CloundBangImages(),
            Announcement(),
        ]

    def test_something(self):
        self.init()
        document = Document()
        for source in self.sources:
            print((json.dumps(source.to_json(), encoding="UTF-8", ensure_ascii=False)))
            data = source.to_json()
            document.add_heading(source._meta.db_table, level=1)
            table = document.add_table(rows=1, cols=4, style='Medium Shading 1')

            cells = table.rows[0].cells

            cells[0].text = "name"
            cells[1].text = "kind"
            cells[2].text = "default"
            cells[3].text = "desc"
            for parameter in data:
                cells = table.add_row().cells
                cells[0].style = 'Medium Grid 1 Accent 1'
                cells[0].text = parameter["name"]
                cells[1].text = parameter["kind"]
                cells[2].text = parameter["default"]
                cells[3].text = parameter["desc"].decode('utf-8')

            document.add_page_break()

            document.save('/Users/fanyangyang/Desktop/demo.docx')


if __name__ == '__main__':
    unittest.main()
